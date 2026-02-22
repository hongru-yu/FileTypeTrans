"""
Word转换器模块的单元测试
"""
import pytest
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
from src.core.converters.docx_converter import DocxConverter


class TestDocxConverter:
    """测试Word转换器类"""

    def test_supports_docx(self):
        """测试是否支持 .docx 文件"""
        converter = DocxConverter()
        assert converter.supports(".docx")

    def test_supports_doc(self):
        """测试是否支持 .doc 文件"""
        converter = DocxConverter()
        assert converter.supports(".doc")

    def test_supports_case_insensitive(self):
        """测试文件扩展名是否大小写不敏感"""
        converter = DocxConverter()
        assert converter.supports(".DOCX")
        assert converter.supports(".DOC")
        assert converter.supports(".Docx")

    def test_does_not_support_other_formats(self):
        """测试不支持其他格式"""
        converter = DocxConverter()
        assert not converter.supports(".pdf")
        assert not converter.supports(".xlsx")
        assert not converter.supports(".txt")

    @patch('src.core.converters.docx_converter.Document')
    def test_convert_docx_to_markdown(self, mock_document_class):
        """测试转换 .docx 文件为 Markdown"""
        # 创建 mock 文档对象
        mock_doc = Mock()
        mock_document_class.return_value = mock_doc

        # 创建 mock 段落
        mock_para1 = Mock()
        mock_para1.text = "标题一"
        mock_para2 = Mock()
        mock_para2.text = "这是第一段内容。"
        mock_para3 = Mock()
        mock_para3.text = "这是第二段内容。"

        mock_doc.paragraphs = [mock_para1, mock_para2, mock_para3]

        # 创建测试文件路径
        test_file = Path("/tmp/test.docx")

        converter = DocxConverter()
        markdown = converter.convert(test_file)

        # 验证输出
        assert "# 原始文件: test.docx" in markdown
        assert "标题一" in markdown
        assert "这是第一段内容。" in markdown
        assert "这是第二段内容。" in markdown

        # 验证 Document 被正确调用
        mock_document_class.assert_called_once_with(test_file)

    @patch('src.core.converters.docx_converter.Document')
    def test_convert_empty_document(self, mock_document_class):
        """测试转换空文档"""
        mock_doc = Mock()
        mock_document_class.return_value = mock_doc
        mock_doc.paragraphs = []

        test_file = Path("/tmp/empty.docx")

        converter = DocxConverter()
        markdown = converter.convert(test_file)

        assert "# 原始文件: empty.docx" in markdown

    @patch('src.core.converters.docx_converter.Document')
    def test_convert_document_with_empty_paragraphs(self, mock_document_class):
        """测试转换包含空段落的文档"""
        mock_doc = Mock()
        mock_document_class.return_value = mock_doc

        mock_para1 = Mock()
        mock_para1.text = "第一段"
        mock_para2 = Mock()
        mock_para2.text = ""
        mock_para3 = Mock()
        mock_para3.text = "第三段"

        mock_doc.paragraphs = [mock_para1, mock_para2, mock_para3]

        test_file = Path("/tmp/mixed.docx")

        converter = DocxConverter()
        markdown = converter.convert(test_file)

        assert "第一段" in markdown
        assert "第三段" in markdown

    @patch('src.core.converters.docx_converter.Document')
    def test_convert_document_with_special_characters(self, mock_document_class):
        """测试转换包含特殊字符的文档"""
        mock_doc = Mock()
        mock_document_class.return_value = mock_doc

        mock_para = Mock()
        mock_para.text = "中文：你好\n数字：123\n符号：@#$%"
        mock_doc.paragraphs = [mock_para]

        test_file = Path("/tmp/special.docx")

        converter = DocxConverter()
        markdown = converter.convert(test_file)

        assert "中文：你好" in markdown
        assert "数字：123" in markdown
        assert "符号：@#$%" in markdown

    @patch('src.core.converters.docx_converter.Document')
    def test_convert_document_with_multiline_paragraphs(self, mock_document_class):
        """测试转换包含多行段落的文档"""
        mock_doc = Mock()
        mock_document_class.return_value = mock_doc

        mock_para1 = Mock()
        mock_para1.text = "第一行\n第二行\n第三行"
        mock_para2 = Mock()
        mock_para2.text = "另一个段落"

        mock_doc.paragraphs = [mock_para1, mock_para2]

        test_file = Path("/tmp/multiline.docx")

        converter = DocxConverter()
        markdown = converter.convert(test_file)

        assert "第一行" in markdown
        assert "第二行" in markdown
        assert "第三行" in markdown
        assert "另一个段落" in markdown

    @patch('src.core.converters.docx_converter.Document')
    def test_convert_large_document(self, mock_document_class):
        """测试转换大文档"""
        mock_doc = Mock()
        mock_document_class.return_value = mock_doc

        # 创建100个段落
        paragraphs = []
        for i in range(100):
            mock_para = Mock()
            mock_para.text = f"段落 {i}"
            paragraphs.append(mock_para)

        mock_doc.paragraphs = paragraphs

        test_file = Path("/tmp/large.docx")

        converter = DocxConverter()
        markdown = converter.convert(test_file)

        assert "段落 0" in markdown
        assert "段落 99" in markdown

    def test_convert_file_not_found(self):
        """测试文件不存在的错误处理"""
        test_file = Path("/tmp/nonexistent.docx")

        converter = DocxConverter()

        with pytest.raises(Exception):
            converter.convert(test_file)

    @patch('src.core.converters.docx_converter.Document')
    def test_convert_corrupted_document(self, mock_document_class):
        """测试损坏文档的错误处理"""
        from docx.opc.exceptions import PackageNotFoundError

        mock_document_class.side_effect = PackageNotFoundError("Document is corrupted")

        test_file = Path("/tmp/corrupted.docx")

        converter = DocxConverter()

        with pytest.raises(PackageNotFoundError):
            converter.convert(test_file)

    def test_init(self):
        """测试初始化"""
        converter = DocxConverter()
        assert converter is not None
        assert hasattr(converter, 'supports')
        assert hasattr(converter, 'convert')
